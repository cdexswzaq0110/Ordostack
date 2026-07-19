"""Build OrdoStack_Interview_Handbook_zh-TW.docx from the markdown source.

The source is a constrained markdown dialect (headings, fenced code, tables,
bullet/numbered lists, blockquotes, **bold**, `inline code`), so this stays a
small deterministic converter instead of a full markdown engine.

    python docs/interview/build_handbook.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "OrdoStack_Interview_Handbook_zh-TW.md"
OUTPUT_PATH = BASE_DIR / "OrdoStack_Interview_Handbook_zh-TW.docx"

BODY_FONT = "Calibri"
EAST_ASIAN_FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def set_run_fonts(run, *, mono: bool = False, size: int | None = None) -> None:
    run.font.name = CODE_FONT if mono else BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), CODE_FONT if mono else BODY_FONT)
    rfonts.set(qn("w:hAnsi"), CODE_FONT if mono else BODY_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIAN_FONT)
    if size is not None:
        run.font.size = Pt(size)


def shade_element(element, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    element.append(shading)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str, *, size: int | None = None) -> None:
    for piece in INLINE_PATTERN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
            set_run_fonts(run, size=size)
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            set_run_fonts(run, mono=True, size=size or 10)
        else:
            run = paragraph.add_run(piece)
            set_run_fonts(run, size=size)


def add_page_number_footer(document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for element_name, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        element = OxmlElement(element_name)
        for key, value in attrs.items():
            element.set(qn(key), value)
        if text is not None:
            element.text = text
        run._element.append(element)
    set_run_fonts(run, size=9)


def add_toc_field(document) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("目錄")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT
    set_run_fonts(run, size=20)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（在 Word 中按右鍵→「更新功能變數」，或 Ctrl+A 後按 F9，即可產生目錄）"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, placeholder, end):
        run._element.append(element)
    set_run_fonts(run)
    document.add_page_break()


def add_cover(document, meta: dict[str, str]) -> None:
    for _ in range(6):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OrdoStack 面試手冊")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = ACCENT
    set_run_fonts(run, size=34)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(meta.get("subtitle", ""))
    run.font.size = Pt(14)
    set_run_fonts(run, size=14)

    document.add_paragraph()
    for line in (
        f"版本：{meta.get('version', '')}（對應 repository 實際程式碼）",
        f"日期：{meta.get('date', '')}",
        "作者：＿＿＿＿＿＿＿＿（請填上你的名字）",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.size = Pt(12)
        set_run_fonts(run, size=12)
    document.add_page_break()


def add_code_block(document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        shade_element(paragraph._p.get_or_add_pPr(), CODE_SHADE)
        run = paragraph.add_run(line if line else " ")
        set_run_fonts(run, mono=True, size=9)


def add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    width = OxmlElement("w:tblW")
    width.set(qn("w:type"), "pct")
    width.set(qn("w:w"), "5000")
    table._tbl.tblPr.append(width)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    table._tbl.tblPr.append(layout)

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            text = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, text, size=9)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                shade_element(cell._tc.get_or_add_tcPr(), "DEEAF6")
    document.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_divider_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells if cell)


def convert(source_path: Path, output_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8")
    meta = dict(re.findall(r"<!-- meta: (\w+)=(.+?) -->", text))
    text = re.sub(r"<!-- meta: .+? -->\n?", "", text)
    lines = text.splitlines()

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIAN_FONT)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.color.rgb = ACCENT
        style.element.get_or_add_rPr()
        rfonts = style.element.rPr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), EAST_ASIAN_FONT)

    add_page_number_footer(document)
    add_cover(document, meta)
    add_toc_field(document)

    index = 0
    seen_first_chapter = False
    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, code_lines)
            continue

        if line.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = parse_table_row(lines[index])
                if not is_divider_row(cells):
                    table_rows.append(cells)
                index += 1
            add_table(document, table_rows)
            continue

        if line.startswith("# "):
            title_text = line[2:].strip()
            if seen_first_chapter:
                document.add_page_break()
            if title_text != "OrdoStack 面試手冊":
                seen_first_chapter = True
                heading = document.add_heading(level=1)
                add_inline_runs(heading, title_text, size=18)
            index += 1
            continue
        if line.startswith("## "):
            heading = document.add_heading(level=2)
            add_inline_runs(heading, line[3:].strip(), size=14)
            index += 1
            continue
        if line.startswith("### "):
            heading = document.add_heading(level=3)
            add_inline_runs(heading, line[4:].strip(), size=12)
            index += 1
            continue

        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_inline_runs(paragraph, line[2:].strip())
            index += 1
            continue
        if re.match(r"^\s*- ", line):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, re.sub(r"^\s*- ", "", line))
            index += 1
            continue
        numbered = re.match(r"^\s*(\d+)\. (.*)", line)
        if numbered:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, f"{numbered.group(1)}. {numbered.group(2)}")
            paragraph.paragraph_format.left_indent = Pt(14)
            index += 1
            continue
        if line.strip() in ("---", "***"):
            index += 1
            continue

        if line.strip():
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, line.strip())
        index += 1

    document.save(output_path)
    print(f"written {output_path} ({output_path.stat().st_size} bytes)")


if __name__ == "__main__":
    convert(SOURCE_PATH, OUTPUT_PATH)
